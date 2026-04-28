"""
Goshwara document generator — produces both .docx and .pdf renditions of the
Bandobast roster with the layout requested by Buldhana District Police:

    • Header: Buldhana District Police + bandobast meta (name, date, time, spot, PS)
    • For each point:
        - Point heading + sector + lat/lng + equipment list
        - Staff table: Sr | Rank | Bakkal | Name | Mobile | Equipment
        - Suchana (instructions) below the staff list
    • Final summary table with rank-wise counts:
        Number of points | Officer (ASP, Dy.SP, PI, API, PSI)
                         | Amaldar (ASI, HC, NPC, PC, LPC)
                         | Female Amaldar | HG | TOTAL
    • Signature block (in-charge officer)

Both DOCX and PDF share the exact same data via _build_data().
"""

from io import BytesIO
from datetime import datetime
import os

# python-docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# reportlab
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether,
    PageBreak,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ---------------------------------------------------------------------------
# Register a font with Devanagari (Marathi) glyph support so the PDF renders
# Marathi text instead of boxes. NotoSans + NotoSansDevanagari are bundled
# from the `fonts-noto-core` Debian package.
# ---------------------------------------------------------------------------

_NOTO_DIR = "/usr/share/fonts/truetype/noto"
_FONT_REGULAR = "Noto"
_FONT_BOLD = "Noto-Bold"
_FONT_ITALIC = "Noto-Italic"

def _try_register_fonts():
    candidates = [
        (_FONT_REGULAR, f"{_NOTO_DIR}/NotoSans-Regular.ttf"),
        (_FONT_BOLD,    f"{_NOTO_DIR}/NotoSans-Bold.ttf"),
        (_FONT_ITALIC,  f"{_NOTO_DIR}/NotoSans-Italic.ttf"),
    ]
    for name, path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
            except Exception:
                pass
    # Devanagari fallback (used by Paragraph thanks to font-family chaining).
    deva = f"{_NOTO_DIR}/NotoSansDevanagari-Regular.ttf"
    if os.path.exists(deva):
        try:
            pdfmetrics.registerFont(TTFont("NotoDeva", deva))
        except Exception:
            pass

_try_register_fonts()


def _font_or(default: str) -> str:
    """Return our Noto font if registered, else the requested default."""
    try:
        return _FONT_REGULAR if _FONT_REGULAR in pdfmetrics.getRegisteredFontNames() else default
    except Exception:
        return default


def _font_bold_or(default: str) -> str:
    try:
        return _FONT_BOLD if _FONT_BOLD in pdfmetrics.getRegisteredFontNames() else default
    except Exception:
        return default


def _font_italic_or(default: str) -> str:
    try:
        return _FONT_ITALIC if _FONT_ITALIC in pdfmetrics.getRegisteredFontNames() else default
    except Exception:
        return default


OFFICER_RANKS = ["ASP", "Dy.SP", "PI", "API", "PSI"]
AMALDAR_RANKS = ["ASI", "HC", "NPC", "PC", "LPC"]


# ---------------------------------------------------------------------------
# Shared data builder
# ---------------------------------------------------------------------------

def _equipment_for(b: dict, point_id: str, staff_id: str) -> str:
    eq_map = (b.get("equipment_assignments") or {}).get(point_id) or {}
    return eq_map.get(staff_id, "") or ""


def _point_rows(b: dict, point: dict, staff_list: list) -> list:
    rows = []
    for i, s in enumerate(staff_list, 1):
        rows.append({
            "sr": i,
            "rank": s.get("rank", ""),
            "bakkal": s.get("bakkal_no", "") or "",
            "name": s.get("name", ""),
            "mobile": s.get("mobile", "") or "-",
            "equipment": _equipment_for(b, point["id"], s["id"]),
        })
    return rows


def _summary_counts(b: dict, point_wise: list) -> dict:
    """
    Return the bottom-of-page summary numbers.
    """
    points_count = len([p for p in (b.get("points") or []) if not p.get("is_reserved")])
    # Walk every allotted staff once
    counts = {r: 0 for r in (OFFICER_RANKS + AMALDAR_RANKS + ["female_amaldar", "hg"])}
    counted = set()
    for entry in point_wise:
        for s in entry["staff"]:
            if s["id"] in counted:
                continue
            counted.add(s["id"])
            stype = s.get("staff_type")
            rank = (s.get("rank") or "").strip()
            gender = (s.get("gender") or "").lower()
            if stype == "officer":
                if rank in OFFICER_RANKS:
                    counts[rank] += 1
            elif stype == "amaldar":
                if gender == "female":
                    counts["female_amaldar"] += 1
                elif rank in AMALDAR_RANKS:
                    counts[rank] += 1
            elif stype == "home_guard":
                counts["hg"] += 1
    total = sum(counts[r] for r in OFFICER_RANKS) + \
            sum(counts[r] for r in AMALDAR_RANKS) + \
            counts["female_amaldar"] + counts["hg"]
    counts["points"] = points_count
    counts["total"] = total
    return counts


# ---------------------------------------------------------------------------
# DOCX generator
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color):
    """Set cell shading."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _bold_cell(cell, text, size=10, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def render_goshwara_docx(b: dict, point_wise: list) -> bytes:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("Buldhana District Police")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Point-wise Goshwara / पॉईंटनिहाय गोषवारा")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Bandobast meta box (single 1-row table)
    meta_text = (
        f"Bandobast: {b.get('name','')}    "
        f"Date: {b.get('date','')}"
        + (f"    Reporting: {b['reporting_time']}" if b.get('reporting_time') else "")
        + f"    Year: {b.get('year','')}    "
        f"Spot: {b.get('spot','-')}    PS: {b.get('ps_name','-')}"
    )
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(meta_text)
    meta_run.font.size = Pt(10)
    meta_run.bold = True
    doc.add_paragraph("")

    # Per-point sections
    for idx, entry in enumerate(point_wise, 1):
        p = entry["point"]
        ps = entry["staff"]

        # Point heading
        ph = doc.add_paragraph()
        run = ph.add_run(
            f"{idx}. {p.get('point_name','')}"
            + ("  (Reserved)" if p.get("is_reserved") else "")
        )
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x31, 0x92)

        # Point meta line
        meta = []
        if p.get("sector"): meta.append(f"Sector: {p['sector']}")
        if p.get("latitude") is not None and p.get("longitude") is not None:
            meta.append(f"Location: {p['latitude']}, {p['longitude']}")
        if p.get("equipment"):
            meta.append(f"Equipment: {', '.join(p['equipment'])}")
        if meta:
            mp = doc.add_paragraph()
            mr = mp.add_run("    ".join(meta))
            mr.font.size = Pt(9)
            mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Staff table
        cols = ["Sr.", "Rank", "Bakkal", "Name", "Mobile", "Equipment"]
        if not ps:
            t = doc.add_table(rows=2, cols=len(cols))
        else:
            t = doc.add_table(rows=len(ps) + 1, cols=len(cols))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Header row
        for ci, name in enumerate(cols):
            cell = t.rows[0].cells[ci]
            _bold_cell(cell, name, size=10, white=True)
            _set_cell_bg(cell, "#2E3192")
        # Data rows
        rows = _point_rows(b, p, ps)
        for ri, r in enumerate(rows, 1):
            vals = [str(r["sr"]), r["rank"], r["bakkal"], r["name"], r["mobile"], r["equipment"] or "-"]
            for ci, v in enumerate(vals):
                c = t.rows[ri].cells[ci]
                c.text = v
                for pp in c.paragraphs:
                    for rr in pp.runs:
                        rr.font.size = Pt(9)
        if not ps:
            c = t.rows[1].cells[0]
            c.merge(t.rows[1].cells[-1])
            c.text = "No staff allotted"
            for pp in c.paragraphs:
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for rr in pp.runs:
                    rr.font.size = Pt(9)
                    rr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Suchana below the staff list
        if p.get("suchana"):
            sp = doc.add_paragraph()
            sr = sp.add_run(f"Suchana / सूचना: {p['suchana']}")
            sr.font.size = Pt(10)
            sr.italic = True
            sr.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
        doc.add_paragraph("")

    # ----- summary table at bottom -----
    counts = _summary_counts(b, point_wise)

    # 13 columns: 1 (points) + 5 officer ranks + 5 amaldar ranks + 1 fem + 1 hg + 1 total
    summary_cols = 1 + len(OFFICER_RANKS) + len(AMALDAR_RANKS) + 3
    sumt = doc.add_table(rows=3, cols=summary_cols)
    sumt.style = "Table Grid"

    # Row 0: Number of points | Officer (merge 5) | Amaldar (merge 5) | Female | HG | TOTAL
    r0 = sumt.rows[0].cells
    _bold_cell(r0[0], "Number of points", size=10)
    _set_cell_bg(r0[0], "#F3F4F6")
    # merge officer cols 1..5
    off_cell = r0[1].merge(r0[1 + len(OFFICER_RANKS) - 1])
    _bold_cell(off_cell, "Officer", size=10)
    _set_cell_bg(off_cell, "#F3F4F6")
    # merge amaldar cols 6..10
    am_start = 1 + len(OFFICER_RANKS)
    am_end = am_start + len(AMALDAR_RANKS) - 1
    am_cell = r0[am_start].merge(r0[am_end])
    _bold_cell(am_cell, "Amaldar", size=10)
    _set_cell_bg(am_cell, "#F3F4F6")
    # remaining
    _bold_cell(r0[am_end + 1], "Female Amaldar", size=10); _set_cell_bg(r0[am_end + 1], "#F3F4F6")
    _bold_cell(r0[am_end + 2], "HG", size=10);             _set_cell_bg(r0[am_end + 2], "#F3F4F6")
    _bold_cell(r0[am_end + 3], "TOTAL", size=10);          _set_cell_bg(r0[am_end + 3], "#F3F4F6")

    # Row 1: blank for points col, then officer rank labels, then amaldar rank labels, then blank for fem/hg/total
    r1 = sumt.rows[1].cells
    r1[0].text = ""
    for i, rk in enumerate(OFFICER_RANKS):
        _bold_cell(r1[1 + i], rk, size=9)
        _set_cell_bg(r1[1 + i], "#F9FAFB")
    for i, rk in enumerate(AMALDAR_RANKS):
        _bold_cell(r1[am_start + i], rk, size=9)
        _set_cell_bg(r1[am_start + i], "#F9FAFB")
    r1[am_end + 1].text = ""
    r1[am_end + 2].text = ""
    r1[am_end + 3].text = ""

    # Row 2: actual counts
    r2 = sumt.rows[2].cells
    r2[0].text = str(counts["points"])
    for i, rk in enumerate(OFFICER_RANKS):
        r2[1 + i].text = str(counts[rk])
    for i, rk in enumerate(AMALDAR_RANKS):
        r2[am_start + i].text = str(counts[rk])
    r2[am_end + 1].text = str(counts["female_amaldar"])
    r2[am_end + 2].text = str(counts["hg"])
    r2[am_end + 3].text = str(counts["total"])
    # center + bigger
    for cell in r2:
        for pp in cell.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rr in pp.runs:
                rr.font.size = Pt(11)
                rr.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Signature block
    sigt = doc.add_table(rows=2, cols=2)
    sigt.alignment = WD_TABLE_ALIGNMENT.CENTER
    sigt.rows[0].cells[0].text = ""
    sigt.rows[0].cells[1].text = ""
    # tall blank cells for signature space
    for cell in sigt.rows[0].cells:
        for pp in cell.paragraphs:
            pp.add_run("\n\n\n")
    # Bottom labels
    left = sigt.rows[1].cells[0]
    right = sigt.rows[1].cells[1]
    lp = left.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run("____________________\nPrepared by")
    lr.font.size = Pt(9)
    rp = right.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    in_charge = b.get("in_charge", "") or "_______________________"
    rr2 = rp.add_run(f"____________________\n{in_charge}\n(In-charge Officer)")
    rr2.bold = True
    rr2.font.size = Pt(10)

    # Footer note
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run(
        f"Generated on {datetime.now().strftime('%d %b %Y %H:%M')}  ·  "
        "Buldhana District Police"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# PDF generator
# ---------------------------------------------------------------------------

NAVY    = colors.HexColor("#2E3192")
SAFFRON = colors.HexColor("#FF9933")
GRAY_50 = colors.HexColor("#F9FAFB")
GRAY_100= colors.HexColor("#F3F4F6")
GRAY_200= colors.HexColor("#E5E7EB")
GRAY_500= colors.HexColor("#6B7280")
GRAY_900= colors.HexColor("#111827")


def render_goshwara_pdf(b: dict, point_wise: list) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=15 * mm, rightMargin=15 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
        title=f"Goshwara - {b.get('name','')}")
    ss = getSampleStyleSheet()
    FN = _font_or("Helvetica")
    FB = _font_bold_or("Helvetica-Bold")
    FI = _font_italic_or("Helvetica-Oblique")
    H1 = ParagraphStyle("H1", parent=ss["Heading1"], fontName=FB,
        fontSize=16, alignment=TA_CENTER, textColor=GRAY_900, spaceAfter=2)
    SUB = ParagraphStyle("Sub", parent=ss["BodyText"], fontName=FN,
        fontSize=10, alignment=TA_CENTER, textColor=GRAY_500, spaceAfter=8)
    META = ParagraphStyle("Meta", parent=ss["BodyText"], fontName=FB,
        fontSize=10, alignment=TA_CENTER, textColor=GRAY_900, spaceAfter=10)
    PT_HEAD = ParagraphStyle("PtHead", parent=ss["Heading2"], fontName=FB,
        fontSize=11.5, textColor=NAVY, spaceBefore=10, spaceAfter=2)
    PT_META = ParagraphStyle("PtMeta", parent=ss["BodyText"], fontName=FN,
        fontSize=8.5, textColor=GRAY_500, spaceAfter=4)
    SUCHANA = ParagraphStyle("Suchana", parent=ss["BodyText"], fontName=FI,
        fontSize=9.5, textColor=colors.HexColor("#B45309"), spaceBefore=4, spaceAfter=10,
        leftIndent=6, backColor=colors.HexColor("#FFF7E6"), borderPadding=4,
        borderColor=SAFFRON, borderWidth=0.5)

    story = []
    story.append(Paragraph("Buldhana District Police", H1))
    story.append(Paragraph("Point-wise Goshwara / पॉईंटनिहाय गोषवारा", SUB))

    meta = (
        f"Bandobast: <b>{b.get('name','')}</b>   ·   "
        f"Date: <b>{b.get('date','')}</b>"
        + (f"   ·   Reporting: <b>{b['reporting_time']}</b>" if b.get('reporting_time') else "")
        + f"   ·   Year: <b>{b.get('year','')}</b>   ·   "
        f"Spot: <b>{b.get('spot','-')}</b>   ·   PS: <b>{b.get('ps_name','-')}</b>"
    )
    story.append(Paragraph(meta, META))

    # Per-point sections
    for idx, entry in enumerate(point_wise, 1):
        p = entry["point"]
        ps = entry["staff"]

        block = []
        title = f"{idx}. {p.get('point_name','')}"
        if p.get("is_reserved"):
            title += " (Reserved)"
        block.append(Paragraph(title, PT_HEAD))

        meta_bits = []
        if p.get("sector"): meta_bits.append(f"Sector: <b>{p['sector']}</b>")
        if p.get("latitude") is not None and p.get("longitude") is not None:
            meta_bits.append(f"Loc: <b>{p['latitude']}, {p['longitude']}</b>")
        if p.get("equipment"):
            meta_bits.append(f"Equipment: <b>{', '.join(p['equipment'])}</b>")
        if meta_bits:
            block.append(Paragraph(" &nbsp;·&nbsp; ".join(meta_bits), PT_META))

        # Staff table
        head = ["Sr.", "Rank", "Bakkal", "Name", "Mobile", "Equipment"]
        rows = _point_rows(b, p, ps)
        if rows:
            data = [head] + [
                [str(r["sr"]), r["rank"], r["bakkal"] or "-", r["name"], r["mobile"], r["equipment"] or "-"]
                for r in rows
            ]
        else:
            data = [head, ["", "", "", "No staff allotted", "", ""]]
        col_widths = [10*mm, 18*mm, 22*mm, 65*mm, 26*mm, 35*mm]
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), FB),
            ("FONTNAME", (0, 1), (-1, -1), FN),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (0, -1), "CENTER"),
            ("ALIGN", (1, 0), (2, -1), "CENTER"),
            ("ALIGN", (4, 0), (5, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.4, GRAY_200),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, GRAY_50]),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        block.append(t)

        # Suchana below the staff list
        if p.get("suchana"):
            block.append(Paragraph(f"<b>Suchana / सूचना:</b> {p['suchana']}", SUCHANA))

        story.append(KeepTogether(block))
        story.append(Spacer(1, 4))

    # --- Summary ---
    counts = _summary_counts(b, point_wise)
    summary_widths = [
        18*mm,  # Number of points
        9*mm, 11*mm, 9*mm, 9*mm, 9*mm,    # Officer ranks (5)
        9*mm, 9*mm, 9*mm, 9*mm, 9*mm,     # Amaldar ranks (5)
        15*mm,                            # Female Amaldar
        9*mm,                             # HG
        12*mm,                            # TOTAL
    ]
    # Row 0 (group headers, with spans)
    row0 = [
        "Number of points",
        "Officer", "", "", "", "",                # spans 5
        "Amaldar", "", "", "", "",                # spans 5
        "Female Amaldar", "HG", "TOTAL",
    ]
    row1 = ["", *OFFICER_RANKS, *AMALDAR_RANKS, "", "", ""]
    row2 = [
        str(counts["points"]),
        *[str(counts[r]) for r in OFFICER_RANKS],
        *[str(counts[r]) for r in AMALDAR_RANKS],
        str(counts["female_amaldar"]),
        str(counts["hg"]),
        str(counts["total"]),
    ]
    sumt = Table([row0, row1, row2], colWidths=summary_widths)
    sumt.setStyle(TableStyle([
        # spans
        ("SPAN", (1, 0), (5, 0)),    # Officer header
        ("SPAN", (6, 0), (10, 0)),   # Amaldar header
        ("SPAN", (0, 0), (0, 1)),    # "Number of points" tall
        ("SPAN", (11, 0), (11, 1)),  # Female amaldar tall
        ("SPAN", (12, 0), (12, 1)),  # HG tall
        ("SPAN", (13, 0), (13, 1)),  # TOTAL tall
        # styling
        ("GRID", (0, 0), (-1, -1), 0.6, GRAY_900),
        ("BACKGROUND", (0, 0), (-1, 1), GRAY_100),
        ("FONTNAME", (0, 0), (-1, 1), FB),
        ("FONTNAME", (0, 2), (-1, 2), FB),
        ("FONTSIZE", (0, 0), (-1, 1), 8.5),
        ("FONTSIZE", (0, 2), (-1, 2), 11),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("<b>Summary</b>", PT_HEAD))
    story.append(sumt)

    # Signature block
    in_charge = b.get("in_charge") or "_______________________"
    sig = Table([
        ["", ""],
        ["____________________", "____________________"],
        ["Prepared by", in_charge],
        ["", "(In-charge Officer)"],
    ], colWidths=[80 * mm, 80 * mm])
    sig.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 1), FN),
        ("FONTNAME", (0, 2), (-1, 3), FB),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("TOPPADDING", (0, 0), (-1, 0), 28),  # signature space
        ("BOTTOMPADDING", (0, 0), (-1, 0), 0),
    ]))
    story.append(Spacer(1, 24))
    story.append(sig)

    # Footer
    foot = ParagraphStyle("F", parent=ss["BodyText"], fontName=FN, fontSize=7.5,
        textColor=GRAY_500, alignment=TA_CENTER, spaceBefore=6)
    story.append(Spacer(1, 16))
    story.append(Paragraph(
        f"Generated on {datetime.now().strftime('%d %b %Y %H:%M')}  ·  "
        "Buldhana District Police",
        foot))

    doc.build(story)
    return buf.getvalue()
